"""
Estudo de Caso 1 - Análise de Texto Científico (VERSÃO MULTI-AGENTE)
- Arquitetura com 3 agentes especializados: Analista Gramatical, Analista de Citações e Analista de Clareza
- Cada agente tem responsabilidade específica no pipeline de análise
- Mantém mesmo resultado final com melhor modularidade
"""

from abc import ABC, abstractmethod
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Protocol, Tuple, Type, Union
import json
from dotenv import load_dotenv
import os
import re
from datetime import datetime

import fitz
import litellm
from crewai import Agent, Crew, LLM, Process, Task
from crewai.tools import BaseTool
from pydantic import BaseModel, Field
from docx import Document

load_dotenv()
# =============================
# TIPOS E CONFIGURAÇÕES (mantidos iguais)
# =============================

class ModelType(Enum):
    LLAMA_8B = "llama_8b"
    LLAMA_70B = "llama_70b"
    CHATGPT = "gpt"


class GravidadeProblema(Enum):
    BAIXA = "baixa"
    MEDIA = "média"
    ALTA = "alta"


@dataclass(frozen=True)
class ConfiguracaoModelo:
    model_name: str
    display_name: str
    api_key: str
    temperature: float = 0.1
    max_tokens: int = 4096
    max_retries: int = 200


@dataclass
class ProblemaEncontrado:
    localizacao: str
    trecho_exato: str
    descricao: str
    sugestao: str
    gravidade: GravidadeProblema
    pagina: Optional[int] = None


@dataclass
class ResultadoAnalise:
    erros_gramaticais: List[ProblemaEncontrado]
    necessidades_citacao: List[ProblemaEncontrado]
    melhorias_clareza: List[ProblemaEncontrado]
    modelo_usado: str
    arquivo_analisado: str
    timestamp: datetime


# =============================
# PROTOCOLS E CLASSES AUXILIARES (mantidos)
# =============================

class LeitorPDF(Protocol):
    def extrair_texto_completo(self, caminho_arquivo: str) -> str: ...
    def extrair_texto_por_paginas(self, caminho_arquivo: str) -> Tuple[str, List[str]]: ...


class ProcessadorJSON(Protocol):
    def extrair_json_de_texto(self, texto: str) -> Optional[Dict]: ...


class ValidadorConteudo(Protocol):
    def validar_problemas_contra_pdf(
        self, problemas: List[Dict], paginas_pdf: List[str]
    ) -> Tuple[List[ProblemaEncontrado], List[Dict]]: ...


class GeradorRelatorio(Protocol):
    def gerar_relatorio_markdown(self, resultado: ResultadoAnalise) -> str: ...


# =============================
# IMPLEMENTAÇÕES CONCRETAS (mantidas iguais)
# =============================

class LeitorDocumentoPyMuPDF:
    """Implementação concreta do leitor de documentos usando PyMuPDF para PDF e python-docx para DOCX"""
    
    def extrair_texto_completo(self, caminho_arquivo: str) -> str:
        if not self._validar_arquivo_existe(caminho_arquivo):
            raise FileNotFoundError(f"Arquivo '{caminho_arquivo}' não encontrado")
        
        try:
            if caminho_arquivo.lower().endswith('.docx'):
                return self._extrair_texto_docx(caminho_arquivo)
            elif caminho_arquivo.lower().endswith('.pdf'):
                return self._extrair_texto_pdf(caminho_arquivo)
            else:
                raise ValueError(f"Formato de arquivo não suportado: {caminho_arquivo}")
                
        except ImportError as e:
            if 'docx' in str(e):
                raise ImportError("python-docx não instalado. Instale com: pip install python-docx")
            else:
                raise ImportError("PyMuPDF não instalado. Instale com: pip install PyMuPDF")
        except Exception as e:
            raise RuntimeError(f"Erro ao ler arquivo '{caminho_arquivo}': {str(e)}")
    
    def extrair_texto_por_paginas(self, caminho_arquivo: str) -> Tuple[str, List[str]]:
        if not self._validar_arquivo_existe(caminho_arquivo):
            raise FileNotFoundError(f"Arquivo '{caminho_arquivo}' não encontrado")
        
        try:
            if caminho_arquivo.lower().endswith('.docx'):
                return self._extrair_texto_docx_por_paragrafos(caminho_arquivo)
            elif caminho_arquivo.lower().endswith('.pdf'):
                return self._extrair_texto_pdf_por_paginas(caminho_arquivo)
            else:
                raise ValueError(f"Formato de arquivo não suportado: {caminho_arquivo}")
        except Exception as e:
            raise RuntimeError(f"Erro ao extrair seções do arquivo '{caminho_arquivo}': {str(e)}")
    
    def _extrair_texto_docx(self, caminho_arquivo: str) -> str:
        doc = Document(caminho_arquivo)
        paragrafos = [paragrafo.text for paragrafo in doc.paragraphs if paragrafo.text.strip()]
        if not paragrafos:
            raise ValueError(f"Arquivo '{caminho_arquivo}' está vazio ou não contém texto extraível")
        return "\n".join(paragrafos)
    
    def _extrair_texto_docx_por_paragrafos(self, caminho_arquivo: str) -> Tuple[str, List[str]]:
        doc = Document(caminho_arquivo)
        paragrafos = [paragrafo.text for paragrafo in doc.paragraphs if paragrafo.text.strip()]
        texto_completo = "\n".join(paragrafos)
        return texto_completo, paragrafos
    
    def _extrair_texto_pdf(self, caminho_arquivo: str) -> str:
        with fitz.open(caminho_arquivo) as documento:
            textos_paginas = [pagina.get_text() for pagina in documento]
            texto_completo = "\n".join(textos_paginas)
            if not texto_completo.strip():
                raise ValueError(f"Arquivo '{caminho_arquivo}' está vazio ou não contém texto extraível")
            return texto_completo
    
    def _extrair_texto_pdf_por_paginas(self, caminho_arquivo: str) -> Tuple[str, List[str]]:
        with fitz.open(caminho_arquivo) as documento:
            paginas = [pagina.get_text() for pagina in documento]
            texto_completo = "\n".join(paginas)
            return texto_completo, paginas
    
    @staticmethod
    def _validar_arquivo_existe(caminho: str) -> bool:
        return os.path.exists(caminho)


class ProcessadorJSONInteligente:
    """Processador de JSON que tenta múltiplas estratégias de extração"""
    
    def extrair_json_de_texto(self, texto: str) -> Optional[Dict]:
        estrategias = [
            self._extrair_json_com_marcadores,
            self._extrair_json_com_codigo_block,
            self._extrair_ultimo_objeto_balanceado,
            self._extrair_json_com_placeholders,
            self._extrair_json_puro
        ]
        
        for estrategia in estrategias:
            resultado = estrategia(texto)
            if resultado:
                return resultado
        return None
    
    def _extrair_json_com_marcadores(self, texto: str) -> Optional[Dict]:
        padrao = r"```json\s*(\{[^}]*\}|\[[^\]]*\])\s*```"
        matches = re.findall(padrao, texto, flags=re.DOTALL | re.IGNORECASE)
        for match in reversed(matches):
            try:
                return json.loads(match)
            except json.JSONDecodeError:
                continue
        return None
    
    def _extrair_json_com_codigo_block(self, texto: str) -> Optional[Dict]:
        padrao = r"```\s*(\{[^}]*\}|\[[^\]]*\])\s*```"
        matches = re.findall(padrao, texto, flags=re.DOTALL)
        for match in reversed(matches):
            try:
                return json.loads(match)
            except json.JSONDecodeError:
                continue
        return None
    
    def _extrair_ultimo_objeto_balanceado(self, texto: str) -> Optional[Dict]:
        ultimo_objeto = None
        pilha_chaves = []
        indice_inicio = None
        
        for indice, caractere in enumerate(texto):
            if caractere == '{':
                if not pilha_chaves:
                    indice_inicio = indice
                pilha_chaves.append('{')
            elif caractere == '}' and pilha_chaves:
                pilha_chaves.pop()
                if not pilha_chaves and indice_inicio is not None:
                    candidato = texto[indice_inicio:indice + 1]
                    ultimo_objeto = candidato
                    indice_inicio = None
        
        return self._tentar_parsear_json(ultimo_objeto)
    
    def _tentar_parsear_json(self, objeto_json: Optional[str]) -> Optional[Dict]:
        if objeto_json:
            try:
                return json.loads(objeto_json)
            except json.JSONDecodeError:
                pass
        return None
    
    def _extrair_json_com_placeholders(self, texto: str) -> Optional[Dict]:
        if "[...]" in texto or "..." in texto:
            padrao_estrutura = r'\{\s*"erros_gramaticais".*?"necessidades_citacao".*?"melhorias_clareza".*?\}'
            match = re.search(padrao_estrutura, texto, re.DOTALL | re.IGNORECASE)
            
            if match:
                json_texto = match.group(0)
                json_texto = re.sub(r'\[\.\.\.?\]', '[]', json_texto)
                json_texto = re.sub(r'\.\.\.', '', json_texto)
                
                try:
                    return json.loads(json_texto)
                except json.JSONDecodeError:
                    return {
                        "erros_gramaticais": [],
                        "necessidades_citacao": [],
                        "melhorias_clareza": []
                    }
        return None

    def _extrair_json_puro(self, texto: str) -> Optional[Dict]:
        try:
            return json.loads(texto.strip())
        except json.JSONDecodeError:
            return None


class ValidadorConteudoPDF:
    """Validador que verifica se problemas encontrados existem no PDF"""
    
    def __init__(self, leitor_pdf: LeitorPDF):
        self._leitor_pdf = leitor_pdf
    
    def validar_problemas_contra_pdf(
        self, problemas: List[Dict], paginas_pdf: List[str]
    ) -> Tuple[List[ProblemaEncontrado], List[Dict]]:
        problemas_validos = []
        problemas_descartados = []
        
        for problema_dict in problemas or []:
            trecho = problema_dict.get("trecho_exato", "").strip()
            
            if not trecho:
                problema_dict["_motivo_descarte"] = "Sem 'trecho_exato'"
                problemas_descartados.append(problema_dict)
                continue
            
            pagina_encontrada = self._localizar_pagina_do_trecho(trecho, paginas_pdf)
            
            if pagina_encontrada is None:
                problema_dict["_motivo_descarte"] = "Trecho não encontrado no PDF"
                problemas_descartados.append(problema_dict)
                continue
            
            problema = self._converter_dict_para_problema(problema_dict, pagina_encontrada)
            problemas_validos.append(problema)
        
        return problemas_validos, problemas_descartados
    
    def _localizar_pagina_do_trecho(self, trecho: str, paginas: List[str]) -> Optional[int]:
        import unicodedata
        from difflib import SequenceMatcher

        def _normalize(text: str) -> str:
            if not text:
                return ""
            text = unicodedata.normalize('NFKD', text)
            text = ''.join(ch for ch in text if not unicodedata.combining(ch))
            text = text.lower()
            text = re.sub(r"[^\w\s]", " ", text)
            text = " ".join(text.split())
            return text

        trecho_normalizado = _normalize(trecho)[:2000]

        if not trecho_normalizado:
            return None

        for indice, texto_pagina in enumerate(paginas):
            pagina_normalizada = _normalize(texto_pagina)
            if trecho_normalizado in pagina_normalizada:
                return indice + 1

        for indice, texto_pagina in enumerate(paginas):
            pagina_normalizada = _normalize(texto_pagina)
            try:
                ratio = SequenceMatcher(None, trecho_normalizado, pagina_normalizada).ratio()
            except Exception:
                ratio = 0.0

            if ratio >= 0.70:
                return indice + 1

            slice_trecho = trecho_normalizado[:200]
            if slice_trecho:
                try:
                    slice_ratio = SequenceMatcher(None, slice_trecho, pagina_normalizada).ratio()
                except Exception:
                    slice_ratio = 0.0

                if slice_ratio >= 0.75:
                    return indice + 1

        return None
    
    def _converter_dict_para_problema(self, problema_dict: Dict, pagina: int) -> ProblemaEncontrado:
        gravidade_str = problema_dict.get("gravidade", "baixa")
        gravidade = GravidadeProblema(gravidade_str)
        
        return ProblemaEncontrado(
            localizacao=problema_dict.get("localizacao", ""),
            trecho_exato=problema_dict.get("trecho_exato", ""),
            descricao=problema_dict.get("descricao", ""),
            sugestao=problema_dict.get("sugestao", ""),
            gravidade=gravidade,
            pagina=pagina
        )


class GeradorRelatorioMarkdown:
    """Gerador de relatórios em formato Markdown"""
    
    def gerar_relatorio_markdown(self, resultado: ResultadoAnalise) -> str:
        linhas = []
        self._adicionar_cabecalho(linhas, resultado)
        self._adicionar_secao_problemas(linhas, "Seção A: Problemas Gramaticais", resultado.erros_gramaticais)
        self._adicionar_secao_problemas(linhas, "Seção B: Necessidades de Citação", resultado.necessidades_citacao)
        self._adicionar_secao_problemas(linhas, "Seção C: Melhorias de Clareza", resultado.melhorias_clareza)
        return "\n".join(linhas)
    
    def _adicionar_cabecalho(self, linhas: List[str], resultado: ResultadoAnalise):
        linhas.extend([
            "# Relatório de Análise Textual Multi-Agente",
            f"**Modelo:** {resultado.modelo_usado}",
            f"**Data:** {resultado.timestamp.strftime('%d/%m/%Y %H:%M:%S')}",
            f"**Arquivo analisado:** {resultado.arquivo_analisado}",
            "",
            "## Resultado Validado (somente itens com evidência literal no documento)",
            "### Processamento Multi-Agente: 3 especialistas trabalharam em paralelo",
            "- **Agente 1**: Analista Gramatical (ortografia, gramática, sintaxe)",
            "- **Agente 2**: Analista de Citações (referenciamento acadêmico)",  
            "- **Agente 3**: Analista de Clareza (legibilidade, coesão textual)",
            ""
        ])
    
    def _adicionar_secao_problemas(self, linhas: List[str], titulo: str, problemas: List[ProblemaEncontrado]):
        linhas.append(f"### {titulo}")
        
        if not problemas:
            linhas.append("_Sem itens validados nessa seção._")
            linhas.append("")
            return
        
        for problema in problemas:
            linhas.extend([
                f"- **Trecho exato:** \"{problema.trecho_exato}\"",
                f"- **Problema:** {problema.descricao}",
                f"- **Sugestão:** {problema.sugestao}",
                f"- **Gravidade:** {problema.gravidade.value}",
                ""
            ])


# =============================
# CONFIGURAÇÕES E FACTORIES
# =============================

class ConfiguradorLiteLLM:
    @staticmethod
    def aplicar_patch_parametros():
        completion_original = litellm.completion
        
        def completion_sem_parametros_problematicos(**parametros):
            parametros_problematicos = ["stop", "stop_sequences", "stops", "stop_tokens"]
            
            for parametro in parametros_problematicos:
                parametros.pop(parametro, None)
            
            if "extra_body" in parametros and isinstance(parametros["extra_body"], dict):
                for parametro in parametros_problematicos:
                    parametros["extra_body"].pop(parametro, None)
            
            return completion_original(**parametros)
        
        litellm.completion = completion_sem_parametros_problematicos


class FabricaModelos:
    """Factory para criação de modelos de IA"""
    
    CONFIGURACOES_PADRAO = {
        ModelType.LLAMA_8B: ConfiguracaoModelo(
            model_name="replicate/meta/meta-llama-3-8b-instruct",
            display_name="Llama 8B",
            api_key=os.getenv('REPLICATE_API_TOKEN')
        ),
        ModelType.LLAMA_70B: ConfiguracaoModelo(
            model_name="replicate/meta/meta-llama-3-70b-instruct",
            display_name="Llama 70B",
            api_key=os.getenv('REPLICATE_API_TOKEN'),
            temperature=0.2,
            max_tokens=8192
        ),
        ModelType.CHATGPT: ConfiguracaoModelo(
            model_name="openai/gpt-4o-mini",
            display_name="ChatGPT 4o Mini",
            api_key=os.getenv('OPENAI_API_KEY')
        )
    }
    
    @classmethod
    def criar_modelo(cls, tipo_modelo: ModelType) -> LLM:
        config = cls.CONFIGURACOES_PADRAO[tipo_modelo]
        if tipo_modelo == ModelType.CHATGPT:
            return LLM(
                model=config.model_name,
                api_key=config.api_key,
                temperature=config.temperature,
                # max_tokens=config.max_tokens,
                max_retries=config.max_retries,
            )
        else: 
            return LLM(
                model=config.model_name,
                api_key=config.api_key,
                drop_params=["stop"],
                temperature=config.temperature,
                # max_tokens=config.max_tokens,
                max_retries=config.max_retries,
            )
    
    @classmethod
    def obter_nome_exibicao(cls, tipo_modelo: ModelType) -> str:
        return cls.CONFIGURACOES_PADRAO[tipo_modelo].display_name


# =============================
# FERRAMENTAS DO CREWAI
# =============================

class ArgumentosLeituraDocumento(BaseModel):
    pdf_path: str = Field(..., description="Caminho do arquivo de documento a ser lido (PDF ou DOCX)")


class FerramentaLeituraDocumento(BaseTool):
    name: str = "LerTextoPDFTool"
    description: str = (
        "Lê o conteúdo de um arquivo de documento (PDF ou DOCX) e retorna o texto extraído. "
        "Use passando um dicionário com a chave 'pdf_path', por ex.: "
        '{"pdf_path": "texto_estudo_caso1.docx"}'
    )
    args_schema: Type[BaseModel] = ArgumentosLeituraDocumento
    
    def __init__(self, leitor_pdf: LeitorPDF):
        super().__init__()
        self._leitor_pdf = leitor_pdf
    
    def _run(self, pdf_path: str) -> str:
        try:
            return self._leitor_pdf.extrair_texto_completo(pdf_path)
        except FileNotFoundError as e:
            return f"Erro: {str(e)}"
        except ValueError as e:
            return f"Aviso: {str(e)}"
        except ImportError as e:
            return f"Erro: {str(e)}"
        except Exception as e:
            return f"Erro ao ler arquivo de documento: {str(e)}"


# =============================
# SISTEMA MULTI-AGENTE (NOVO)
# =============================

class FabricaAgentesEspecializados:
    """Factory para criar os 3 agentes especializados"""
    
    @staticmethod
    def criar_agente_gramatical(modelo: LLM, ferramenta_documento: FerramentaLeituraDocumento) -> Agent:
        """Agente 1: Especialista em análise gramatical"""
        return Agent(
            role="Analista Gramatical Especializado",
            goal=(
                "Identificar problemas de ortografia, gramática e sintaxe em textos acadêmicos "
                "com foco em correções que melhorem a qualidade técnica do documento."
            ),
            backstory=(
                "Especialista em correção de textos com 15 anos de experiência em revisão "
                "de documentos acadêmicos. Expert em detectar erros ortográficos, problemas "
                "de concordância, regência e sintaxe. Foco na precisão técnica."
            ),
            verbose=False,
            tools=[ferramenta_documento],
            llm=modelo,
            max_iter=2,
            max_execution_time=300,
        )
    
    @staticmethod
    def criar_agente_citacoes(modelo: LLM, ferramenta_documento: FerramentaLeituraDocumento) -> Agent:
        """Agente 2: Especialista em análise de citações"""
        return Agent(
            role="Analista de Citações Acadêmicas",
            goal=(
                "Identificar afirmações, dados e claims que necessitam de citações ou "
                "referências para manter o rigor acadêmico do documento."
            ),
            backstory=(
                "Doutor em Metodologia Científica com expertise em normas acadêmicas. "
                "Especialista em identificar quando afirmações precisam de fundamentação "
                "bibliográfica. Conhece padrões ABNT, APA e outros sistemas de citação."
            ),
            verbose=False,
            tools=[ferramenta_documento],
            llm=modelo,
            max_iter=2,
            max_execution_time=300,
        )
    
    @staticmethod
    def criar_agente_clareza(modelo: LLM, ferramenta_documento: FerramentaLeituraDocumento) -> Agent:
        """Agente 3: Especialista em clareza e estilo"""
        return Agent(
            role="Analista de Clareza e Estilo",
            goal=(
                "Identificar problemas de clareza, coesão textual e legibilidade que "
                "comprometem a compreensão e fluidez do documento acadêmico."
            ),
            backstory=(
                "Editor científico com vasta experiência em melhorar a legibilidade "
                "de textos acadêmicos. Expert em detectar ambiguidades, frases mal "
                "estruturadas e problemas de coesão que dificultam a compreensão."
            ),
            verbose=False,
            tools=[ferramenta_documento],
            llm=modelo,
            max_iter=2,
            max_execution_time=300,
        )


# Substitua a classe FabricaTarefasEspecializadas pelos prompts otimizados:

class FabricaTarefasEspecializadas:
    """Factory para criar as 3 tarefas especializadas com prompts otimizados"""
    
    @staticmethod
    def criar_tarefa_gramatical(agente: Agent, ferramenta_documento: FerramentaLeituraDocumento, caminho_arquivo: str, texto_documento: str) -> Task:
        """Tarefa 1: Análise gramatical especializada - PROMPT OTIMIZADO"""
        
        return Task(
            description=f"""
**MISSÃO**: Você é um corretor especializado em gramática da língua portuguesa. Sua ÚNICA responsabilidade é identificar erros técnicos da língua.

**PASSO 1 - LEITURA OBRIGATÓRIA**:
Use a ferramenta LerTextoPDFTool para ler o documento:
{{"pdf_path": "{caminho_arquivo}"}}

**PASSO 2 - ANÁLISE FOCADA**:
Identifique APENAS estes tipos de erro:

🔍 **ORTOGRAFIA**:
- Palavras escritas incorretamente (ex: "questãos" → "questões")
- Acentuação inadequada
- Uso incorreto de hífen

🔍 **CONCORDÂNCIA**:
- Verbal: sujeito-verbo (ex: "Os dados mostra" → "Os dados mostram")
- Nominal: substantivo-adjetivo (ex: "questões específico" → "questões específicas")

🔍 **REGÊNCIA**:
- Verbal: preposições com verbos (ex: "assistir o filme" → "assistir ao filme")
- Nominal: preposições com substantivos

🔍 **SINTAXE E PONTUAÇÃO**:
- Vírgulas mal posicionadas
- Pontos e vírgulas inadequados
- Estrutura frasal problemática

**PASSO 3 - SAÍDA ESTRUTURADA**:
```json
{{
  "erros_gramaticais": [
    {{
      "localizacao": "Seção 3.1, segundo parágrafo",
      "trecho_exato": "texto que está LITERALMENTE no documento",
      "descricao": "Erro de concordância verbal",
      "sugestao": "correção específica proposta",
      "gravidade": "média"
    }}
  ]
}}
```

**VALIDAÇÃO CRÍTICA**:
✅ "trecho_exato" deve existir LITERALMENTE no documento
✅ Se não encontrar erros gramaticais, retorne: {{"erros_gramaticais": []}}
❌ NÃO inclua problemas de citação, clareza ou conteúdo
❌ NÃO invente trechos que não existem

**TEXTO DE BACKUP** (use apenas se a ferramenta falhar):
{texto_documento[:12000]}{"...[TRUNCADO]" if len(texto_documento) > 12000 else ""}
""",
            expected_output="JSON válido contendo lista de problemas gramaticais com localização precisa",
            agent=agente,
            tools=[ferramenta_documento]
        )
    
    @staticmethod
    def criar_tarefa_citacoes(agente: Agent, ferramenta_documento: FerramentaLeituraDocumento, caminho_arquivo: str, texto_documento: str) -> Task:
        """Tarefa 2: Análise de necessidades de citação - PROMPT OTIMIZADO"""
        
        return Task(
            description=f"""
**MISSÃO**: Você é um especialista em normas acadêmicas. Sua ÚNICA responsabilidade é identificar afirmações que precisam de fundamentação bibliográfica.

**PASSO 1 - LEITURA OBRIGATÓRIA**:
Use a ferramenta LerTextoPDFTool para ler o documento:
{{"pdf_path": "{caminho_arquivo}"}}

**PASSO 2 - IDENTIFICAÇÃO CRÍTICA**:
Procure APENAS estes casos que EXIGEM citação:

📚 **DADOS ESPECÍFICOS**:
- Estatísticas sem fonte (ex: "85% dos usuários relataram...")
- Números precisos de pesquisas
- Porcentagens de estudos

📚 **AFIRMAÇÕES CATEGÓRICAS**:
- Claims definitivos sobre fenômenos (ex: "A IA sempre produz resultados melhores")
- Declarações sobre causas e efeitos específicos
- Comparações entre métodos/tecnologias

📚 **TEORIAS E CONCEITOS**:
- Menção a teorias específicas sem autor
- Definições técnicas sem referência
- Modelos ou frameworks conceituais

📚 **RESULTADOS DE PESQUISA**:
- "Estudos mostram que..."
- "Pesquisas indicam..."
- "Foi comprovado que..."

**O QUE NÃO PRECISA DE CITAÇÃO**:
❌ Conhecimento geral (ex: "A internet é amplamente utilizada")
❌ Definições básicas universalmente aceitas
❌ Descrição de metodologia própria do autor

**PASSO 3 - SAÍDA ESTRUTURADA**:
```json
{{
  "necessidades_citacao": [
    {{
      "localizacao": "Seção 2, primeiro parágrafo",
      "trecho_exato": "texto que está LITERALMENTE no documento",
      "descricao": "Afirmação estatística sem fonte comprobatória",
      "sugestao": "Adicionar citação de pesquisa que comprove esta estatística",
      "gravidade": "alta"
    }}
  ]
}}
```

**VALIDAÇÃO RIGOROSA**:
✅ Apenas afirmações que REALMENTE precisam de fundamentação
✅ "trecho_exato" deve existir LITERALMENTE no documento
✅ Se não encontrar necessidades de citação, retorne: {{"necessidades_citacao": []}}
❌ NÃO inclua problemas gramaticais ou de clareza
❌ NÃO seja excessivamente rigoroso com conhecimento básico

**TEXTO DE BACKUP** (use apenas se a ferramenta falhar):
{texto_documento[:12000]}{"...[TRUNCADO]" if len(texto_documento) > 12000 else ""}
""",
            expected_output="JSON válido contendo lista de necessidades de citação com justificativa clara",
            agent=agente,
            tools=[ferramenta_documento]
        )
    
    @staticmethod
    def criar_tarefa_clareza(agente: Agent, ferramenta_documento: FerramentaLeituraDocumento, caminho_arquivo: str, texto_documento: str) -> Task:
        """Tarefa 3: Análise de clareza e estilo - PROMPT OTIMIZADO"""
        
        return Task(
            description=f"""
**MISSÃO**: Você é um editor científico especializado em legibilidade. Sua ÚNICA responsabilidade é identificar problemas que comprometem a compreensão do texto.

**PASSO 1 - LEITURA OBRIGATÓRIA**:
Use a ferramenta LerTextoPDFTool para ler o documento:
{{"pdf_path": "{caminho_arquivo}"}}

**PASSO 2 - ANÁLISE DE LEGIBILIDADE**:
Identifique APENAS estes problemas de clareza:

✨ **AMBIGUIDADE**:
- Frases com múltiplas interpretações possíveis
- Pronomes com referência unclear (ex: "isso", "aquilo" sem antecedente claro)
- Modificadores mal posicionados

✨ **COMPLEXIDADE EXCESSIVA**:
- Períodos muito longos (>3 linhas) sem pausas adequadas
- Estruturas sintáticas desnecessariamente complexas
- Acúmulo excessivo de subordinadas

✨ **COESÃO TEXTUAL**:
- Transições abruptas entre ideias
- Parágrafos sem conectores lógicos
- Sequência de ideias confusa

✨ **JARGÃO TÉCNICO**:
- Termos técnicos sem explicação no primeiro uso
- Acronimos não expandidos
- Linguagem excessivamente hermética para o público-alvo

✨ **FLUXO DE LEITURA**:
- Repetições desnecessárias de palavras próximas
- Cacofonia ou aliterações problemáticas
- Ritmo de leitura prejudicado

**O QUE NÃO É PROBLEMA DE CLAREZA**:
❌ Erros gramaticais (outro agente cuida)
❌ Falta de citações (outro agente cuida)
❌ Escolhas estilísticas legítimas do autor

**PASSO 3 - SAÍDA ESTRUTURADA**:
```json
{{
  "melhorias_clareza": [
    {{
      "localizacao": "Seção 3.2, terceiro parágrafo",
      "trecho_exato": "texto que está LITERALMENTE no documento",
      "descricao": "Período excessivamente longo que dificulta compreensão",
      "sugestao": "Dividir em duas frases para melhorar legibilidade",
      "gravidade": "média"
    }}
  ]
}}
```

**CRITÉRIOS DE QUALIDADE**:
✅ Foque em melhorias que REALMENTE aumentem a legibilidade
✅ Preserve o tom acadêmico nas sugestões
✅ "trecho_exato" deve existir LITERALMENTE no documento
✅ Se não encontrar problemas de clareza, retorne: {{"melhorias_clareza": []}}
❌ NÃO sugira mudanças que alterem o significado
❌ NÃO inclua problemas que são de outros agentes

**TEXTO DE BACKUP** (use apenas se a ferramenta falhar):
{texto_documento[:12000]}{"...[TRUNCADO]" if len(texto_documento) > 12000 else ""}
""",
            expected_output="JSON válido contendo lista de melhorias de clareza com sugestões construtivas",
            agent=agente,
            tools=[ferramenta_documento]
        )

class ConsolidadorResultados:
    """Consolida resultados dos 3 agentes especializados"""
    
    @staticmethod
    def consolidar_analises(
        resultado_gramatical: str,
        resultado_citacoes: str, 
        resultado_clareza: str,
        processador_json: ProcessadorJSON
    ) -> Dict:
        """Consolida os 3 resultados em uma estrutura unificada"""
        
        # Extrair JSONs de cada agente
        json_gramatical = processador_json.extrair_json_de_texto(resultado_gramatical)
        json_citacoes = processador_json.extrair_json_de_texto(resultado_citacoes)
        json_clareza = processador_json.extrair_json_de_texto(resultado_clareza)
        
        # Estrutura consolidada
        resultado_consolidado = {
            "erros_gramaticais": json_gramatical.get("erros_gramaticais", []) if json_gramatical else [],
            "necessidades_citacao": json_citacoes.get("necessidades_citacao", []) if json_citacoes else [],
            "melhorias_clareza": json_clareza.get("melhorias_clareza", []) if json_clareza else []
        }
        
        return resultado_consolidado


# =============================
# SERVIÇO PRINCIPAL MULTI-AGENTE
# =============================

class ServicoAnaliseTextoMultiAgente:
    """Serviço principal com arquitetura multi-agente"""
    
    def __init__(
        self,
        leitor_pdf: LeitorPDF,
        processador_json: ProcessadorJSON,
        validador_conteudo: ValidadorConteudo,
        gerador_relatorio: GeradorRelatorio
    ):
        self._leitor_pdf = leitor_pdf
        self._processador_json = processador_json
        self._validador_conteudo = validador_conteudo
        self._gerador_relatorio = gerador_relatorio
    
    def analisar_documento_multiagente(
        self, 
        tipo_modelo: ModelType, 
        caminho_arquivo: str
    ) -> Tuple[Optional[ResultadoAnalise], Optional[str]]:
        """Executa análise com 3 agentes especializados"""
        try:
            # Validações iniciais
            if not os.path.exists(caminho_arquivo):
                return None, f"Arquivo '{caminho_arquivo}' não encontrado!"
            
            # Configurar modelo e infraestrutura
            modelo = FabricaModelos.criar_modelo(tipo_modelo)
            nome_modelo = FabricaModelos.obter_nome_exibicao(tipo_modelo)
            
            ferramenta_documento = FerramentaLeituraDocumento(self._leitor_pdf)
            
            # Extrair texto do documento como fallback
            try:
                texto_documento = self._leitor_pdf.extrair_texto_completo(caminho_arquivo)
            except Exception:
                texto_documento = ""
            
            print(f"🤖 Criando 3 agentes especializados para {nome_modelo}...")
            
            # Criar os 3 agentes especializados
            agente_gramatical = FabricaAgentesEspecializados.criar_agente_gramatical(modelo, ferramenta_documento)
            agente_citacoes = FabricaAgentesEspecializados.criar_agente_citacoes(modelo, ferramenta_documento)
            agente_clareza = FabricaAgentesEspecializados.criar_agente_clareza(modelo, ferramenta_documento)
            
            print("📝 ETAPA 1: Análise gramatical...")
            
            # Execução do Agente 1: Análise Gramatical
            tarefa_gramatical = FabricaTarefasEspecializadas.criar_tarefa_gramatical(
                agente_gramatical, ferramenta_documento, caminho_arquivo, texto_documento
            )
            
            crew_gramatical = Crew(
                agents=[agente_gramatical],
                tasks=[tarefa_gramatical],
                process=Process.sequential,
                verbose=False,
                max_execution_time=600,
            )
            
            resultado_gramatical = str(crew_gramatical.kickoff()).strip()
            
            print("📚 ETAPA 2: Análise de citações...")
            
            # Execução do Agente 2: Análise de Citações
            tarefa_citacoes = FabricaTarefasEspecializadas.criar_tarefa_citacoes(
                agente_citacoes, ferramenta_documento, caminho_arquivo, texto_documento
            )
            
            crew_citacoes = Crew(
                agents=[agente_citacoes],
                tasks=[tarefa_citacoes],
                process=Process.sequential,
                verbose=False,
                max_execution_time=600,
            )
            
            resultado_citacoes = str(crew_citacoes.kickoff()).strip()
            
            print("✨ ETAPA 3: Análise de clareza...")
            
            # Execução do Agente 3: Análise de Clareza
            tarefa_clareza = FabricaTarefasEspecializadas.criar_tarefa_clareza(
                agente_clareza, ferramenta_documento, caminho_arquivo, texto_documento
            )
            
            crew_clareza = Crew(
                agents=[agente_clareza],
                tasks=[tarefa_clareza],
                process=Process.sequential,
                verbose=False,
                max_execution_time=600,
            )
            
            resultado_clareza = str(crew_clareza.kickoff()).strip()
            
            print("🔄 ETAPA 4: Consolidando resultados...")
            
            # Consolidar resultados dos 3 agentes
            json_consolidado = ConsolidadorResultados.consolidar_analises(
                resultado_gramatical, resultado_citacoes, resultado_clareza, self._processador_json
            )
            
            if not json_consolidado:
                return None, f"Não foi possível consolidar resultados de {nome_modelo}"
            
            # Validar contra documento original
            _, paginas_pdf = self._leitor_pdf.extrair_texto_por_paginas(caminho_arquivo)
            
            erros_validos, _ = self._validador_conteudo.validar_problemas_contra_pdf(
                json_consolidado.get("erros_gramaticais", []), paginas_pdf
            )
            citacoes_validas, _ = self._validador_conteudo.validar_problemas_contra_pdf(
                json_consolidado.get("necessidades_citacao", []), paginas_pdf
            )
            clareza_valida, _ = self._validador_conteudo.validar_problemas_contra_pdf(
                json_consolidado.get("melhorias_clareza", []), paginas_pdf
            )
            
            # Criar resultado final
            resultado_final = ResultadoAnalise(
                erros_gramaticais=erros_validos,
                necessidades_citacao=citacoes_validas,
                melhorias_clareza=clareza_valida,
                modelo_usado=f"{nome_modelo} (Multi-Agente)",
                arquivo_analisado=caminho_arquivo,
                timestamp=datetime.now()
            )
            
            print(f"✅ {nome_modelo}: Análise multi-agente concluída")
            print(f"📊 Problemas encontrados: {len(erros_validos)} gramáticais, {len(citacoes_validas)} citações, {len(clareza_valida)} clareza")
            
            return resultado_final, None
            
        except Exception as e:
            erro = f"Erro na análise multi-agente com {nome_modelo}: {str(e)}"
            print(f"❌ {erro}")
            return None, erro


# =============================
# GERENCIADOR MULTI-AGENTE
# =============================

class GerenciadorEstudoCasoMultiAgente:
    """Gerenciador adaptado para arquitetura multi-agente"""
    
    def __init__(self, servico_analise: ServicoAnaliseTextoMultiAgente, gerador_relatorio: GeradorRelatorio):
        self._servico_analise = servico_analise
        self._gerador_relatorio = gerador_relatorio
        self._diretorio_resultados = Path("resultados_estudo_caso")
    
    def executar_estudo_completo(self, caminho_arquivo: str = "texto_estudo_caso1.docx"):
        print("🚀 ESTUDO DE CASO 1 - ANÁLISE DE TEXTO CIENTÍFICO (MULTI-AGENTE)")
        print("🤖 Arquitetura: 3 Agentes Especializados")
        print("   • Agente 1: Analista Gramatical (ortografia, gramática, sintaxe)")
        print("   • Agente 2: Analista de Citações (referenciamento acadêmico)")
        print("   • Agente 3: Analista de Clareza (legibilidade, coesão textual)")
        print("🎯 Pipeline: Análise Paralela → Consolidação → Validação\n")
        
        self._garantir_diretorio_resultados()
        
        resultados_sucesso = {}
        
        for tipo_modelo in ModelType:
            nome_modelo = FabricaModelos.obter_nome_exibicao(tipo_modelo)
            print(f"\n{'='*60}\n🤖 PROCESSAMENTO MULTI-AGENTE: {nome_modelo}\n{'='*60}")
            
            # Usar método multi-agente
            resultado, erro = self._servico_analise.analisar_documento_multiagente(tipo_modelo, caminho_arquivo)

            if resultado:
                caminho_relatorio = self._salvar_relatorio_individual(resultado, tipo_modelo)
                resultados_sucesso[tipo_modelo] = {
                    "resultado": resultado,
                    "arquivo": caminho_relatorio,
                    "status": "✅ Sucesso"
                }
                print(f"✅ {nome_modelo}: Concluído com sucesso")
            else:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                nome_arquivo = f"relatorio_{tipo_modelo.value}_{timestamp}_FAILED.md"
                caminho_arquivo_saida = self._diretorio_resultados / nome_arquivo

                linhas = [
                    f"# Relatório (FALHA) - {nome_modelo} Multi-Agente",
                    "",
                    f"**Arquivo analisado:** {caminho_arquivo}",
                    f"**Data:** {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}",
                    "",
                    "## Erro",
                    "",
                    f"{erro}",
                    "",
                    "## Arquitetura Multi-Agente",
                    "",
                    "- 3 agentes especializados trabalhando em paralelo",
                    "- Pipeline: Gramatical → Citações → Clareza → Consolidação",
                ]

                with open(caminho_arquivo_saida, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(linhas))

                resultados_sucesso[tipo_modelo] = {
                    "resultado": None,
                    "arquivo": str(caminho_arquivo_saida),
                    "status": f"❌ Falhou: {erro}"
                }

                print(f"❌ {nome_modelo}: {erro}")
        
        print("\n🎉 ESTUDO DE CASO 1 MULTI-AGENTE CONCLUÍDO!")
        print(f"📁 Resultados salvos em: {self._diretorio_resultados}/")
        sucessos = sum(1 for r in resultados_sucesso.values() if r["status"].startswith("✅"))
        print(f"📊 Taxa de sucesso: {sucessos}/{len(ModelType)} modelos")
        print(f"🤖 Cada modelo utilizou 3 agentes especializados trabalhando em paralelo")
    
    def _garantir_diretorio_resultados(self):
        self._diretorio_resultados.mkdir(parents=True, exist_ok=True)
    
    def _salvar_relatorio_individual(self, resultado: ResultadoAnalise, tipo_modelo: ModelType) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f"relatorio_{tipo_modelo.value}_{timestamp}.md"
        caminho_arquivo = self._diretorio_resultados / nome_arquivo
        
        conteudo_relatorio = self._gerador_relatorio.gerar_relatorio_markdown(resultado)
        
        with open(caminho_arquivo, "w", encoding="utf-8") as arquivo:
            arquivo.write(conteudo_relatorio)
        
        print(f"📄 Relatório multi-agente salvo: {caminho_arquivo}")
        return str(caminho_arquivo)


# =============================
# PONTO DE ENTRADA PRINCIPAL
# =============================

def main():
    """Função principal do programa"""
    ConfiguradorLiteLLM.aplicar_patch_parametros()
    
    # Criar dependências
    leitor_pdf = LeitorDocumentoPyMuPDF()
    processador_json = ProcessadorJSONInteligente()
    validador_conteudo = ValidadorConteudoPDF(leitor_pdf)
    gerador_relatorio = GeradorRelatorioMarkdown()
    
    # Criar serviço multi-agente
    servico_analise = ServicoAnaliseTextoMultiAgente(
        leitor_pdf=leitor_pdf,
        processador_json=processador_json,
        validador_conteudo=validador_conteudo,
        gerador_relatorio=gerador_relatorio
    )
    
    # Criar gerenciador multi-agente e executar
    gerenciador = GerenciadorEstudoCasoMultiAgente(servico_analise, gerador_relatorio)
    gerenciador.executar_estudo_completo('arquivo_estudo_1/texto_estudo_caso1.docx')


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n🛑 Execução interrompida pelo usuário. Encerrando com segurança...")